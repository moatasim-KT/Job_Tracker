from flask import Blueprint, request, jsonify, flash
import os
import glob
from job_tracker.utils.llm_parser import JobDescriptionParser

cover_letter_bp = Blueprint('cover_letter', __name__)
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.dirname(__file__)), '..', 'uploads')
COVER_LETTERS_DIR = os.path.join(UPLOAD_FOLDER, 'cover_letters')
OTHER_DOCS_DIR = os.path.join(UPLOAD_FOLDER, 'other_docs')

@cover_letter_bp.route('/cover-letter/generate', methods=['POST'])
def generate_cover_letter():
    data = request.get_json()
    job_description = data.get('job_description', '')
    # Find the uploaded CV (if any)
    cv_text = ''
    for ext in ['pdf', 'doc', 'docx', 'txt']:
        path = os.path.join(UPLOAD_FOLDER, f'cv.{ext}')
        if os.path.exists(path):
            try:
                if ext == 'pdf':
                    import PyPDF2
                    with open(path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        cv_text = "\n".join(page.extract_text() or '' for page in reader.pages)
                elif ext in ['doc', 'docx']:
                    import docx
                    doc = docx.Document(path)
                    cv_text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    with open(path, 'r', encoding='utf-8') as f:
                        cv_text = f.read()
            except Exception as e:
                return jsonify({'error': f'Could not read CV: {str(e)}'}), 400
            break
    if not cv_text:
        return jsonify({'error': 'No CV uploaded or unable to read CV.'}), 400
    if not job_description:
        return jsonify({'error': 'Job description is required.'}), 400

    # Gather all additional docs (cover letters + other docs)
    def extract_text_from_file(path):
        ext = path.split('.')[-1].lower()
        try:
            if ext == 'pdf':
                import PyPDF2
                with open(path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return "\n".join(page.extract_text() or '' for page in reader.pages)
            elif ext in ['doc', 'docx']:
                import docx
                doc = docx.Document(path)
                return "\n".join([para.text for para in doc.paragraphs])
            else:
                with open(path, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception:
            return ''

    additional_docs = []
    # Cover letters
    for fname in glob.glob(os.path.join(COVER_LETTERS_DIR, '*')):
        additional_docs.append(extract_text_from_file(fname))
    # Other docs
    for fname in glob.glob(os.path.join(OTHER_DOCS_DIR, '*')):
        additional_docs.append(extract_text_from_file(fname))

    try:
        letter = JobDescriptionParser.generate_cover_letter(job_description, cv_text, additional_docs)
        # If the letter looks like an error message, surface it as an error
        if letter.startswith('[LLM API Error') or letter.startswith('[Could not extract cover letter'):
            return jsonify({'error': letter}), 500
        return jsonify({'cover_letter': letter})
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        return jsonify({'error': f'{str(e)}\n{tb}'}), 500
